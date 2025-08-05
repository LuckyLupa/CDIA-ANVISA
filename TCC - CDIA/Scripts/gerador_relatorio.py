from docx import Document

def gerar_relatorio(dados, template_path, output_path):
    # Carregar template
    doc = Document(template_path)
    
    # Substituir placeholders simples nos parágrafos
    for p in doc.paragraphs:
        for key, value in dados.items():
            if f"{{{{{key}}}}}" in p.text:
                p.text = p.text.replace(f"{{{{{key}}}}}", value)
    
    # Substituir placeholders nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in dados.items():
                    if f"{{{{{key}}}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{{{key}}}}}", value)

    # Salvar arquivo final
    doc.save(output_path)
    print(f"Relatório gerado com sucesso: {output_path}")

if __name__ == "__main__":
    # Dados simulados da fatura piloto (RIOgaleão TIC Julho/2025)
    dados_fatura = {
        "NUM_FATURA": "8000475862",
        "DATA_EMISSAO": "25/06/2025",
        "DATA_VENCIMENTO": "20/07/2025",
        "PERIODO": "Julho/2025",
        "CNPJ_CONCESSIONARIA": "19.726.111/0001-08",
        "NUM_CONTRATO": "2204/2019",
        "VALOR_CUSTO_TI": "11.179,62",
        "VALOR_CONSUMO_TI": "169,99",
        "VALOR_TOTAL_BRUTO": "11.349,61",
        "NUM_EMPENHO": "2025NE000123",
        "ID_CONTRATOS_GOV": "123456",
        "LINK_FATURA": "https://sei.anvisa.gov.br/fatura/8000475862.pdf",
        "HASH_PDF": "a3f5c8d9b4e8f72c4569e1a3d8c9f02aab89c1f234abcd56789ef01234567890"
    }

    template_path = "templates/modelo_relatorio_v2.docx"
    output_path = "data/exemplos_saida/relatorio_aux_riogaleao_tic_2025-07.docx"
    
    gerar_relatorio(dados_fatura, template_path, output_path)
